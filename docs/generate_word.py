"""
将 Markdown 格式的《系统实现与测试报告》转换为 Word 文档
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import re

doc = Document()

# ========== 样式设置 ==========
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 标题样式
for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = '微软雅黑'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if level == 1:
        heading_style.font.size = Pt(22)
        heading_style.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)
    elif level == 2:
        heading_style.font.size = Pt(16)
        heading_style.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)
    elif level == 3:
        heading_style.font.size = Pt(13)
        heading_style.font.color.rgb = RGBColor(0x3A, 0x6E, 0x9F)

# ========== 封面 ==========
for _ in range(6):
    doc.add_paragraph()

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run('系统实现与测试报告')
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)

doc.add_paragraph()

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle_para.add_run('SukinOS - 浏览器端桌面操作系统')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

for _ in range(4):
    doc.add_paragraph()

info_items = [
    ('项目名称', 'SukinOS'),
    ('版本号', 'v0.1.0'),
    ('报告日期', '2025年6月16日'),
    ('文档状态', 'V1.0'),
]
for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}：{value}')
    run.font.size = Pt(14)

doc.add_page_break()

# ========== 目录页 ==========
doc.add_heading('目录', level=1)
toc_items = [
    '一、系统实现',
    '    1.1 开发环境',
    '    1.2 项目架构概述',
    '    1.3 核心模块实现说明',
    '    1.4 数据库表结构',
    '    1.5 运行截图',
    '二、系统测试',
    '    2.1 单元测试用例表及运行结果',
    '    2.2 集成测试用例表',
    '    2.3 系统测试结果',
    '    2.4 缺陷统计与分析',
    '    2.5 测试结论',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ========== 正文内容 ==========

def add_code_block(doc, code_text, language=''):
    """添加代码块（灰底等宽字体）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # 添加灰色背景
    pPr = p._p.get_or_add_pPr()
    shading = pPr.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F5F5F5',
        qn('w:val'): 'clear'
    })
    pPr.append(shading)
    for line in code_text.split('\n'):
        if line.strip():
            run = p.add_run(line + '\n')
        else:
            run = p.add_run('\n')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_table(doc, headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    table.autofit = True
    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()

# ========== 一、系统实现 ==========
doc.add_heading('一、系统实现', level=1)

doc.add_heading('1.1 开发环境', level=2)

doc.add_heading('1.1.1 前端开发环境', level=3)
add_table(doc,
    ['项目', '说明'],
    [
        ['操作系统', 'Windows 11'],
        ['运行时', 'Node.js 20+'],
        ['构建工具', 'Vite 7'],
        ['框架', 'React 19'],
        ['状态管理', 'Redux Toolkit + redux-persist'],
        ['路由', 'react-router-dom v7'],
        ['样式方案', 'CSS Modules + Tailwind CSS 4 + Emotion'],
        ['编辑器', 'Monaco Editor'],
        ['拖拽', '@dnd-kit/core'],
        ['动画', 'GSAP'],
        ['图表', 'ECharts 6'],
        ['UI库', 'MUI v7'],
        ['编译', 'Babel standalone (浏览器端JSX转译)'],
    ]
)

doc.add_heading('1.1.2 后端开发环境', level=3)
add_table(doc,
    ['项目', '说明'],
    [
        ['操作系统', 'Windows 11 / Linux (systemd)'],
        ['运行时', 'Python 3.11.5'],
        ['Web框架', 'FastAPI 0.137.1'],
        ['ORM', 'SQLAlchemy 2.0.41'],
        ['数据库', 'MySQL 8.0'],
        ['缓存', 'Redis 6.x'],
        ['异步任务', 'Celery (Redis broker)'],
        ['认证', 'JWT (HS256) + bcrypt'],
        ['测试框架', 'pytest 9.1.0'],
    ]
)

doc.add_heading('1.2 项目架构概述', level=2)
doc.add_paragraph(
    'SukinOS 是一个运行在浏览器端的桌面操作系统 Web 应用，采用前后端分离架构。'
    '前端基于 React 19 + Redux Toolkit，模拟桌面操作系统的核心体验——窗口化多任务、'
    '应用动态加载、虚拟文件系统、进程间通信及安全沙箱隔离。数据完全存储在用户浏览器本地（IndexedDB）。'
    '后端基于 FastAPI + SQLAlchemy + Celery + Redis，提供用户认证管理、SukinOS 应用商店、'
    '请求日志审计、WebSocket 实时通信等服务。'
)

doc.add_paragraph(
    '系统整体架构分为三层：浏览器端（SukinOS 前端 React 应用，包含 DeskBook 桌面环境、'
    'Resources 应用层、Process Kernel 进程内核、VFS 虚拟文件系统、Redux Store、IndexedDB 存储），'
    '网络层（HTTP/WebSocket 通信），服务器端（FastAPI 后端，包含 Auth 认证、User API、'
    'SukinOS App API、WebSocket 实时通信、System Config 配置管理、Request Log 请求日志、'
    'Middleware Stack 中间件栈，连接 MySQL、Redis、Celery 等基础设施）。'
)

doc.add_heading('1.3 核心模块实现说明', level=2)

doc.add_heading('1.3.1 后端核心模块', level=3)

doc.add_heading('(1) 应用启动与初始化 (main.py + exit.py)', level=3)
doc.add_paragraph(
    '系统采用多进程架构，主进程负责启动 FastAPI 和 Celery Worker 两个子进程。'
    'exit.py 中的 init_config() 负责注册全局异常处理器、添加中间件栈、挂载静态文件目录、'
    '注册路由模块、初始化数据库表并清除僵尸在线状态。'
)

doc.add_heading('(2) JWT 双 Token 认证系统 (Api/Common/Auth.py)', level=3)
doc.add_paragraph(
    '系统实现了完整的 Access Token + Refresh Token 双 Token 机制。'
    'Access Token 有效期 30 分钟用于业务请求，Refresh Token 有效期 7 天用于静默刷新。'
    '支持"记住我"模式、bcrypt 密码哈希存储、三级权限模型（角色/位运算/超级管理员）。'
    'verify_auth() 工厂函数实现了四阶段认证流程：校验 Access Token → 静默刷新 → 用户状态校验 → 详细权限校验。'
)
add_code_block(doc, '''def create_access_token(data: dict, remember: bool = True):
    """创建短期 Access Token (用于业务请求)"""
    to_encode = data.copy()
    expires_delta = timedelta(seconds=ACCESS_TOKEN_EXPIRE_MINUTES)
    expire = datetime.now(timezone.utc) + expires_delta
    to_encode.update({
        "exp": int(expire.timestamp()),
        "type": "access",
        "remember": remember
    })
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt, int(expires_delta.total_seconds())''')

doc.add_heading('(3) 用户管理 API (Api/User/userMain.py)', level=3)
doc.add_paragraph('提供完整的用户生命周期管理接口，包括验证码获取/校验、登录（账号密码/验证码两种方式）、Token 刷新/校验、在线状态变更、注销、资料修改（含头像上传）、密码修改等功能。')

doc.add_heading('(4) SukinOS 应用商店 (Api/SukinOS/appManage.py)', level=3)
doc.add_paragraph(
    '采用策略模式设计存储层（BaseStorage 抽象基类 + LocalStorage 本地实现），'
    '支持未来扩展云存储。提供应用上传（含磁盘空间检查和延迟等待重试）、版本更新（乐观锁版本号自动递增）、'
    '公开应用列表分页查询、批量检查更新、搜索、我的上传记录查询、删除等功能。'
)

doc.add_heading('(5) WebSocket 实时通信 (Api/SocketApi/loginDate.py)', level=3)
doc.add_paragraph(
    'ConnectionManager 使用 Dict[int, Set[WebSocket]] 支持同一用户多标签页连接。'
    '实现心跳检测（35秒超时）、登录/登出时间精确追踪（跨天切分）、'
    '时段活跃度统计（凌晨/上午/下午/晚上四时段）、周度时长汇总。'
)

doc.add_heading('(6) 请求日志中间件 (Middleware/requestLog.py)', level=3)
doc.add_paragraph(
    '纯 ASGI 中间件，实现完整的请求审计日志。支持敏感信息脱敏（password、token、cookie、secret 等 12 个敏感键自动替换为 ***）、'
    'JSON 载荷采样压缩（列表超过 3 个元素时截断）、请求体大小限制（超过 12000 字符自动截断）、'
    'multipart/form-data 和二进制流自动省略。'
)

doc.add_heading('(7) 统一响应格式 (Api/Re_Rs.py)', level=3)
doc.add_paragraph(
    '实现 api_response() 统一响应函数，自动将下划线命名转为驼峰命名（递归转换）、'
    '支持二进制响应和自定义媒体类型、统一的 JSON 响应格式 {code, message, data}。'
)

doc.add_heading('1.3.2 前端核心模块', level=3)
doc.add_paragraph('前端核心模块包括：')
doc.add_paragraph('• 应用入口 (src/main.jsx)：挂载 React 根节点，包裹 Redux Provider + PersistGate + RouterProvider', style='List Bullet')
doc.add_paragraph('• Redux Store 动态架构 (src/store/main.jsx)：支持运行时动态注入 Reducer 和中间件', style='List Bullet')
doc.add_paragraph('• 进程内核 Kernel (src/sukinos/utils/process/kernel.js)：类微内核架构，单例 Kernel 拆分为 10 个子模块（Core/Cache/Flags/Instance/Internals/Lifecycle/Messaging/Registry/ResourceAccess/Settings）', style='List Bullet')
doc.add_paragraph('• 虚拟文件系统 VFS (src/sukinos/utils/file/fileKernel.js)：基于 IndexedDB 的类 Unix 文件系统，inodeMap + treeMap 双索引结构', style='List Bullet')
doc.add_paragraph('• 安全沙箱 (src/sukinos/utils/security.js)：CDN 白名单、Blob 递归沙箱、Storage 代理、SDK 冻结、安全 Fetch 代理', style='List Bullet')
doc.add_paragraph('• 10 个内置系统应用：开发者中心、文件管理器、笔记本、设置、开始菜单、应用商店、本地开发、系统管理、画板、表格', style='List Bullet')

doc.add_heading('1.4 数据库表结构', level=2)
doc.add_paragraph('系统使用 MySQL 8.0，共 8 张核心表：')

add_table(doc,
    ['序号', '表名', '模型类', '用途'],
    [
        ['1', 'users', 'User', '用户账户（账号、密码、邮箱、权限JSON、头像等）'],
        ['2', 'user_expand', 'UserExpand', '用户扩展信息（在线状态、累计在线时长）'],
        ['3', 'user_time_behavior', 'UserTimeBehavior', '用户时间行为（每日在线、时段活跃统计）'],
        ['4', 'user_time_week', 'TimeWeek', '周度时长汇总'],
        ['5', 'sukinos_app', 'D_SukinosApp', 'SukinOS 应用商店（文件、版本、状态、乐观锁）'],
        ['6', 'system_configs', 'D_SystemConfig', '系统配置键值存储'],
        ['7', 'system_updates', 'D_SystemUpdate', '系统更新日志'],
        ['8', 'request_logs', 'D_RequestLog', '请求审计日志'],
    ]
)

doc.add_heading('1.5 运行截图', level=2)
doc.add_paragraph(
    '前端运行截图位于 docs/screenshot/ 目录（8 张），包括系统桌面、应用窗口、文件管理器、开发者中心等界面。'
    '架构图位于 docs/diagrams/ 目录（16 张），包括 API 相关、APP 相关、内核相关、系统相关的架构图。'
)

doc.add_page_break()

# ========== 二、系统测试 ==========
doc.add_heading('二、系统测试', level=1)

doc.add_heading('2.1 单元测试用例表及运行结果', level=2)

doc.add_paragraph(
    '测试框架：pytest 9.1.0。测试文件位置：b:/myfastapi_/tests/。'
    '共 5 个测试文件，108 个测试用例。'
)

add_table(doc,
    ['测试模块', '测试文件', '用例数', '通过', '跳过', '失败', '通过率'],
    [
        ['JWT认证与密码', 'test_auth.py', '23', '22', '1', '0', '95.7%'],
        ['数据模型', 'test_models.py', '22', '22', '0', '0', '100%'],
        ['日期工具函数', 'test_date_utils.py', '14', '14', '0', '0', '100%'],
        ['统一响应工具', 'test_re_rs.py', '22', '22', '0', '0', '100%'],
        ['安全工具函数', 'test_security.py', '27', '27', '0', '0', '100%'],
        ['总计', '', '108', '107', '1', '0', '99.1%'],
    ]
)

doc.add_heading('2.1.1 JWT认证与密码测试用例 (test_auth.py)', level=3)
add_table(doc,
    ['编号', '测试用例', '测试内容', '结果'],
    [
        ['AUTH-01', 'test_hash_password', 'bcrypt 密码哈希生成', '✅ 通过'],
        ['AUTH-02', 'test_verify_correct', '验证正确密码', '✅ 通过'],
        ['AUTH-03', 'test_verify_wrong', '验证错误密码', '✅ 通过'],
        ['AUTH-04', 'test_verify_empty', '验证空密码', '✅ 通过'],
        ['AUTH-05', 'test_different_salts', '相同密码不同哈希', '✅ 通过'],
        ['AUTH-06', 'test_various_lengths', '不同长度密码', '✅ 通过'],
        ['AUTH-07', 'test_gensalt_rounds', '自定义轮数', '✅ 通过'],
        ['AUTH-08', 'test_pwd_context_hash', 'passlib 兼容性', '⏭ 跳过'],
        ['AUTH-09~15', 'Token创建测试(7项)', 'Access/Refresh Token', '✅ 通过'],
        ['AUTH-16~20', 'Token解码测试(5项)', '解码/错误密钥/过期', '✅ 通过'],
        ['AUTH-21~23', '时长配置测试(3项)', '30分钟/7天/比较', '✅ 通过'],
    ]
)

doc.add_heading('2.1.2 数据模型测试用例 (test_models.py)', level=3)
add_table(doc,
    ['编号', '测试用例', '测试内容', '结果'],
    [
        ['MOD-01~06', 'User模型测试(6项)', '创建/权限/关联/root/头像/行为', '✅ 通过'],
        ['MOD-07~10', 'App模型测试(4项)', '创建/默认版本/状态/私密性', '✅ 通过'],
        ['MOD-11~15', '版本号递增测试(5项)', '递增/进位/三段/None/错误格式', '✅ 通过'],
        ['MOD-16~19', '请求日志测试(4项)', '创建/错误/请求体/查询参数', '✅ 通过'],
        ['MOD-20~21', '系统配置/更新测试(2项)', '配置创建/更新日志', '✅ 通过'],
        ['MOD-22', '表结构完整性', '8张表全部创建', '✅ 通过'],
    ]
)

doc.add_heading('2.1.3 日期工具函数测试用例 (test_date_utils.py)', level=3)
add_table(doc,
    ['编号', '测试用例', '测试内容', '结果'],
    [
        ['DT-01~05', '时段映射测试(5项)', '凌晨/上午/下午/晚上/边界', '✅ 通过'],
        ['DT-06', '星期映射完整性', '7天全部映射', '✅ 通过'],
        ['DT-07~09', '周起始时间戳(3项)', '类型/周一/非未来', '✅ 通过'],
        ['DT-10~12', '同周判断(3项)', '同天/同周不同天/不同周', '✅ 通过'],
    ]
)

doc.add_heading('2.1.4 安全工具函数测试用例 (test_security.py)', level=3)
add_table(doc,
    ['编号', '测试用例', '测试内容', '结果'],
    [
        ['SEC-01~05', '文本截断测试(5项)', '短/长/None/边界/空', '✅ 通过'],
        ['SEC-06~15', 'JSON脱敏测试(10项)', '敏感字段/嵌套/大小写/列表', '✅ 通过'],
        ['SEC-16~20', 'JSON解析测试(5项)', '有效/无效/嵌套/数组', '✅ 通过'],
        ['SEC-21~25', '请求体解码(5项)', 'JSON/multipart/二进制', '✅ 通过'],
        ['SEC-26~27', '响应体解码(2项)', 'JSON/二进制', '✅ 通过'],
    ]
)

doc.add_heading('2.1.5 测试运行结果', level=3)
add_code_block(doc, '''======================= 107 passed, 1 skipped in 3.24s =======================''')

doc.add_heading('2.2 集成测试用例表', level=2)
add_table(doc,
    ['编号', '测试场景', '测试接口', '预期输出'],
    [
        ['INT-01', '用户注册-登录-Token刷新', 'POST /user/status/token → /refresh/token', '200 + 双Token'],
        ['INT-02', '验证码登录流程', 'GET/POST 验证码 → 登录', '200 + Token'],
        ['INT-03', '修改密码流程', 'POST /user/status/password/update', '200 + 新Token'],
        ['INT-04', '应用上传-审核-上架', 'POST /upload → /updateStatus', '状态变更'],
        ['INT-05', '应用更新版本流程', 'POST /sukinos/app/update', '版本号递增'],
        ['INT-06', 'WebSocket在线追踪', 'ws://host/ws/date/{id}', '在线状态更新'],
        ['INT-07', '请求日志记录', '任意HTTP请求', '日志写入DB'],
        ['INT-08', '系统健康检查', 'GET /system/status', 'DB/Redis/资源状态'],
        ['INT-09', '文件上传（头像）', 'POST /profile/update', '头像URL更新'],
        ['INT-10', '应用搜索与分页', 'GET /searchApp + /appList', '分页结果'],
    ]
)

doc.add_heading('2.3 系统测试结果', level=2)

doc.add_heading('2.3.1 功能测试结果', level=3)
add_table(doc,
    ['功能模块', '测试项', '通过', '状态'],
    [
        ['用户认证', '登录/注册/Token刷新/登出', '4/4', '✅ 正常'],
        ['密码管理', '修改密码/验证码验证', '2/2', '✅ 正常'],
        ['个人信息', '资料修改/头像上传', '2/2', '✅ 正常'],
        ['应用商店', '上传/更新/列表/搜索/删除', '5/5', '✅ 正常'],
        ['应用审核', '状态变更/强制删除/统计', '3/3', '✅ 正常'],
        ['在线状态', 'WebSocket连接/心跳/断开', '3/3', '✅ 正常'],
        ['请求日志', '记录/脱敏/截断/查询', '4/4', '✅ 正常'],
        ['系统管理', '配置管理/用户管理/健康检查', '3/3', '✅ 正常'],
        ['总计', '', '26/26 (100%)', '✅ 全部正常'],
    ]
)

doc.add_heading('2.4 缺陷统计与分析', level=2)

add_table(doc,
    ['严重级别', '数量', '描述', '状态'],
    [
        ['高', '1', 'passlib 与 bcrypt 5.0.0 不兼容', '已识别，可降级bcrypt'],
        ['中', '1', 'SQLite不支持MEDIUMTEXT类型（仅影响测试）', '已绕过'],
        ['低', '2', '部分模块缺少类型注解；test.py为空', '建议补充'],
    ]
)

doc.add_paragraph(
    '根因分析：bcrypt 5.0.0 移除了 __about__ 属性，passlib 的版本检测代码依赖此属性。'
    'MEDIUMTEXT 是 MySQL 专有类型，SQLite 不支持，单元测试通过 SQLite 兼容模型类绕过。'
)

doc.add_heading('2.5 测试结论', level=2)

doc.add_paragraph(
    '本次测试共编写 5 个测试文件，覆盖 6 个核心模块，共计 108 个测试用例。'
    '通过 107 个（99.1%），跳过 1 个（0.9%，因依赖兼容性），失败 0 个。'
)

doc.add_paragraph('覆盖的核心功能：')
doc.add_paragraph('✅ JWT Token 创建与验证（HS256 算法）', style='List Bullet')
doc.add_paragraph('✅ bcrypt 密码哈希与验证', style='List Bullet')
doc.add_paragraph('✅ 数据库模型完整性（8张表全部验证）', style='List Bullet')
doc.add_paragraph('✅ 版本号递增逻辑', style='List Bullet')
doc.add_paragraph('✅ 用户在线时间追踪（时段映射、周计算）', style='List Bullet')
doc.add_paragraph('✅ 统一响应格式（驼峰转换、状态码映射）', style='List Bullet')
doc.add_paragraph('✅ 请求日志安全（敏感信息脱敏、文本截断、JSON安全解析）', style='List Bullet')
doc.add_paragraph('✅ 请求/响应体解码（JSON/multipart/二进制）', style='List Bullet')

doc.add_heading('质量评估', level=3)
add_table(doc,
    ['维度', '评分', '说明'],
    [
        ['功能完整性', '⭐⭐⭐⭐⭐', '核心功能全部实现并通过测试'],
        ['代码规范性', '⭐⭐⭐⭐', '遵循PEP8，有清晰的模块划分'],
        ['安全性', '⭐⭐⭐⭐⭐', 'JWT双Token + bcrypt + 敏感信息脱敏 + 三级权限'],
        ['可维护性', '⭐⭐⭐⭐', '模块化设计，策略模式，清晰目录结构'],
        ['测试覆盖', '⭐⭐⭐⭐', '核心模块单元测试覆盖良好'],
    ]
)

doc.add_paragraph(
    '结论：SukinOS 系统后端核心模块经过单元测试验证，功能正确、安全可靠。'
    'JWT 双 Token 认证机制、bcrypt 密码哈希、请求日志脱敏、统一响应格式、数据模型等核心功能均通过测试。'
    '系统架构设计合理，模块划分清晰，具备良好的可扩展性和可维护性。'
)

# ========== 保存 ==========
output_path = 'b:/sukinOs/docs/系统实现与测试报告.docx'
doc.save(output_path)
print(f'Word 文档已保存至: {output_path}')
